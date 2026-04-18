# -*- coding: utf-8 -*-
"""PART 1 – KOIB RAG Preprocessing (extract text, OCR, figures, models)"""

# ============================================================================
# 1. Установка и импорт зависимостей (для Colab)
# ============================================================================
!pip uninstall -q -y langchain langchain-core langchain-community langchain-text-splitters langchain-huggingface 2>/dev/null || true
!pip install -q pymupdf python-docx Pillow numpy tqdm pytesseract easyocr sentence-transformers
!apt-get install -qq tesseract-ocr tesseract-ocr-rus -y

import os
import re
import json
import csv
import io
import time
import datetime
import hashlib
from pathlib import Path
from collections import defaultdict
import logging

import fitz               # pymupdf
from docx import Document as DocxDocument
from PIL import Image
import pytesseract
import numpy as np
from tqdm import tqdm

# EasyOCR (ленивая загрузка)
_easyocr_reader = None
def get_easyocr_reader():
    global _easyocr_reader
    if _easyocr_reader is None:
        import easyocr
        _easyocr_reader = easyocr.Reader(['ru', 'en'], gpu=True, verbose=False)
    return _easyocr_reader

# ============================================================================
# 2. Монтирование Google Drive (для Colab)
# ============================================================================
from google.colab import drive
DRIVE_MOUNT = '/content/drive'
if not os.path.isdir(os.path.join(DRIVE_MOUNT, 'MyDrive')):
    drive.mount(DRIVE_MOUNT)
else:
    print("✅ Google Drive already mounted")

# ============================================================================
# 3. Константы и пути
# ============================================================================
DOCS_DIR = Path("/content/drive/MyDrive/Koib/docs")
OUTPUT_DIR = Path("/content/drive/MyDrive/Koib/koib_rag_GLM1")

CLASSIFIED_DIR = OUTPUT_DIR / "classified"
OCR_RESULTS_DIR = OUTPUT_DIR / "ocr_results"
FIGURES_DIR = OUTPUT_DIR / "figures"
METADATA_DIR = OUTPUT_DIR / "metadata"
LOGS_DIR = OUTPUT_DIR / "logs"

for d in [OUTPUT_DIR, CLASSIFIED_DIR, OCR_RESULTS_DIR, FIGURES_DIR,
          METADATA_DIR, LOGS_DIR]:
    d.mkdir(parents=True, exist_ok=True)

# Параметры обработки
OCR_DPI = 300
OCR_MIN_TEXT_CHARS = 50
MIN_IMAGE_WIDTH = 80
MIN_IMAGE_HEIGHT = 80
SCREENSHOT_AREA_THRESHOLD = 0.80
TEXT_DENSITY_THRESHOLD = 0.35

# Паттерны моделей КОИБ
KOIB_MODEL_PATTERNS = {
    "koib2010": [
        r"КОИБ[-\s]?2010", r"КОИБ\s*2010", r"0912054",
        r"PRINT_KOIB2010", r"2010.*руководство", r"руководство.*2010",
        r"модель\s*17404049\.438900\.001",
    ],
    "koib2017a": [
        r"КОИБ[-\s]?2017\s*[АA]", r"КОИБ[-\s]?2017А",
        r"модель\s*17404049\.5013009\.008-01", r"17404049\.5013009",
        r"PRINT_KOIB2017[АA]", r"2017[АA].*руководство",
    ],
    "koib2017b": [
        r"КОИБ[-\s]?2017\s*[БB]", r"КОИБ[-\s]?2017Б",
        r"БАВУ\.201119", r"0912053", r"PRINT_KOIB2017[БB]",
        r"2017[БB].*руководство",
    ],
}

MODEL_DISPLAY_NAMES = {
    "koib2010": "КОИБ-2010",
    "koib2017a": "КОИБ-2017А",
    "koib2017b": "КОИБ-2017Б",
    "unknown": "Неизвестная модель",
}

FIGURE_CAPTION_PATTERNS = [
    re.compile(r"(Рис(?:ун(?:ок|ке))[\s.]?\s*[\d.]+[^\n]*)", re.IGNORECASE),
    re.compile(r"(Рис\.?\s*[\d.]+[^\n]*)", re.IGNORECASE),
    re.compile(r"(Фиг\.?\s*[\d.]+[^\n]*)", re.IGNORECASE),
    re.compile(r"(Рисунок\s+\d+[\.\s][^\n]*)", re.IGNORECASE),
]

# ============================================================================
# 4. Утилиты: очистка текста, хеши, определение модели
# ============================================================================
def clean_text(text):
    if not text:
        return ""
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[^\x20-\x7E\u0400-\u04FF\u2116\n\r\t]', '', text)
    lines = [line.strip() for line in text.split('\n')]
    return '\n'.join(lines).strip()

def text_hash(text):
    return hashlib.md5(text.encode('utf-8', errors='ignore')).hexdigest()[:12]

def normalize_model_key(key):
    key = str(key).strip().lower()
    return key if key in KOIB_MODEL_PATTERNS else "unknown"

def detect_model_in_text(text):
    if not text or len(text.strip()) < 5:
        return ("unknown", 0.0)
    text_upper = text.upper()
    scores = {}
    for model_key, patterns in KOIB_MODEL_PATTERNS.items():
        match_count = 0
        total_matches = 0
        for pat in patterns:
            matches = re.findall(pat, text, re.IGNORECASE)
            if matches:
                match_count += 1
                total_matches += len(matches)
        if match_count > 0:
            scores[model_key] = match_count * 10 + total_matches
    if not scores:
        return ("unknown", 0.0)
    best_model = max(scores, key=scores.get)
    best_score = scores[best_model]
    confidence = min(best_score / 30.0, 1.0)
    return (best_model, round(confidence, 3))

def detect_model_from_filename(filename):
    fn = filename.lower()
    for model_key, patterns in KOIB_MODEL_PATTERNS.items():
        for pat in patterns:
            if re.search(pat, fn, re.IGNORECASE):
                return model_key
    return "unknown"

def find_figure_caption(text, max_distance=300):
    if not text:
        return ""
    for pat in FIGURE_CAPTION_PATTERNS:
        match = pat.search(text)
        if match:
            caption = match.group(1).strip()
            if len(caption) > 3:
                return caption
    return ""

# ============================================================================
# 5. OCR функции (pytesseract + easyocr fallback)
# ============================================================================
def ocr_image(image_pil, lang='rus+eng'):
    if image_pil is None:
        return ""
    # pytesseract
    try:
        text_tess = pytesseract.image_to_string(image_pil, lang=lang, config='--psm 6')
        text_tess = clean_text(text_tess)
        if len(text_tess) >= 30:
            return text_tess
    except:
        text_tess = ""
    # easyocr fallback
    try:
        reader = get_easyocr_reader()
        img_np = np.array(image_pil)
        results = reader.readtext(img_np, paragraph=True, detail=0)
        text_easy = clean_text('\n'.join(results))
        if len(text_easy) >= 20:
            return text_easy
    except:
        text_easy = ""
    return text_tess if len(text_tess) >= len(text_easy) else text_easy

def detect_scanned_page(page, min_text_chars=50):
    try:
        text = page.get_text("text").strip()
        images = page.get_images(full=True)
        page_rect = page.rect
        page_area = page_rect.width * page_rect.height
        if len(text) < min_text_chars:
            for img_info in images:
                try:
                    xref = img_info[0]
                    base_image = page.parent.extract_image(xref)
                    if base_image:
                        img = Image.open(io.BytesIO(base_image["image"]))
                        w, h = img.size
                        if w * h > page_area * 0.3:
                            return True
                except:
                    continue
            return len(text) < 10
    except:
        return True
    return False

def ocr_pdf_page(page, dpi=300, save_result=True, page_num=0, source_file=""):
    try:
        pix = page.get_pixmap(dpi=dpi)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        ocr_text = ocr_image(img)
        if save_result and ocr_text:
            safe_name = re.sub(r'[^\w\-.]', '_', source_file)
            ocr_path = OCR_RESULTS_DIR / f"{safe_name}_page{page_num:03d}.txt"
            ocr_path.write_text(ocr_text, encoding='utf-8')
            img_path = OCR_RESULTS_DIR / f"{safe_name}_page{page_num:03d}.png"
            img.save(img_path, "PNG")
        return (ocr_text, len(ocr_text) > 20)
    except:
        return ("", False)

def is_page_screenshot(img_bytes, min_size=(800, 1000)):
    try:
        img = Image.open(io.BytesIO(img_bytes))
        w, h = img.size
        aspect = w / h if h > 0 else 0
        is_large = w >= min_size[0] and h >= min_size[1]
        is_a4_like = 0.5 < aspect < 0.85
        return is_large and is_a4_like
    except:
        return False

def compute_text_density(img_pil):
    try:
        img_np = np.array(img_pil)
        reader = get_easyocr_reader()
        results = reader.readtext(img_np, paragraph=False)
        if not results:
            return 0.0
        total_text_area = 0
        img_area = img_pil.width * img_pil.height
        for (bbox, text, conf) in results:
            x_coords = [p[0] for p in bbox]
            y_coords = [p[1] for p in bbox]
            box_area = (max(x_coords) - min(x_coords)) * (max(y_coords) - min(y_coords))
            total_text_area += box_area
        return total_text_area / img_area if img_area > 0 else 0.0
    except:
        return 0.0

# ============================================================================
# 6. Обработчик PDF
# ============================================================================
class PdfProcessor:
    def __init__(self, path, source_model=""):
        self.path = Path(path)
        self.filename = self.path.name
        self.source_model = normalize_model_key(source_model)
        self.text_blocks = []
        self.figure_candidates = []
        self.ocr_page_count = 0
        self.processing_time = 0

    def process(self):
        start_time = time.time()
        try:
            doc = fitz.open(str(self.path))
        except Exception as e:
            print(f"  ❌ Error opening {self.filename}: {e}")
            return ([], [])

        total_pages = len(doc)
        print(f"  📄 {self.filename}: {total_pages} pages")

        # Определение модели по первым 5 страницам
        if self.source_model == "unknown":
            sample_text = ""
            for pn in range(min(5, total_pages)):
                try:
                    sample_text += doc[pn].get_text("text") + "\n"
                except:
                    pass
            detected_model, conf = detect_model_in_text(sample_text)
            if conf > 0.1:
                self.source_model = detected_model
                print(f"    🏷️ Detected model: {MODEL_DISPLAY_NAMES.get(self.source_model, self.source_model)} (conf={conf})")

        for page_num in range(total_pages):
            try:
                page = doc[page_num]
                self._process_page(page, page_num+1, doc)
            except Exception as e:
                print(f"    ⚠️ Page {page_num+1} error: {e}")

        doc.close()
        self.processing_time = time.time() - start_time
        print(f"    ✅ Extracted {len(self.text_blocks)} blocks, {len(self.figure_candidates)} figures, OCR pages: {self.ocr_page_count} ({self.processing_time:.1f}s)")
        return (self.text_blocks, self.figure_candidates)

    def _process_page(self, page, page_num, doc):
        is_scanned = detect_scanned_page(page, OCR_MIN_TEXT_CHARS)
        if is_scanned:
            ocr_text, success = ocr_pdf_page(page, dpi=OCR_DPI, save_result=True,
                                              page_num=page_num, source_file=self.filename)
            if ocr_text:
                self.text_blocks.append({
                    "text": ocr_text,
                    "block_type": "ocr_text",
                    "page": page_num,
                    "model": self.source_model,
                    "source": self.filename,
                })
                self.ocr_page_count += 1
        else:
            text = clean_text(page.get_text("text"))
            if text:
                blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
                headings = self._extract_headings(blocks, page_num)
                self.text_blocks.append({
                    "text": text,
                    "block_type": "text",
                    "page": page_num,
                    "model": self.source_model,
                    "source": self.filename,
                    "headings": headings,
                })
            self._extract_images(page, page_num, text)

    def _extract_headings(self, blocks, page_num):
        headings = []
        font_sizes = []
        for block in blocks:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    size = span.get("size", 0)
                    if size > 0:
                        font_sizes.append(size)
        if not font_sizes:
            return headings
        median_size = sorted(font_sizes)[len(font_sizes)//2]
        heading_threshold = median_size * 1.3
        for block in blocks:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    size = span.get("size", 0)
                    text = span.get("text", "").strip()
                    if size >= heading_threshold and len(text) > 3 and len(text) < 200:
                        if text not in headings:
                            headings.append(text)
        return headings

    def _extract_images(self, page, page_num, surrounding_text=""):
        try:
            images = page.get_images(full=True)
            page_rect = page.rect
            page_area = page_rect.width * page_rect.height
        except:
            return
        caption = find_figure_caption(surrounding_text)
        for img_idx, img_info in enumerate(images):
            try:
                xref = img_info[0]
                base_image = page.parent.extract_image(xref)
                if not base_image:
                    continue
                img_bytes = base_image["image"]
                img_ext = base_image.get("ext", "png")
                img = Image.open(io.BytesIO(img_bytes))
                w, h = img.size
                if w < MIN_IMAGE_WIDTH or h < MIN_IMAGE_HEIGHT:
                    continue
                img_area = w * h
                if img_area > page_area * SCREENSHOT_AREA_THRESHOLD:
                    continue
                if img_area > 50000:
                    td = compute_text_density(img)
                    if td > TEXT_DENSITY_THRESHOLD:
                        continue
                fig_id = f"{self.filename}_p{page_num}_img{img_idx}"
                self.figure_candidates.append({
                    "fig_id": fig_id,
                    "source": self.filename,
                    "page": page_num,
                    "model": self.source_model,
                    "width": w,
                    "height": h,
                    "ext": img_ext,
                    "caption": caption,
                    "img_bytes": img_bytes,
                    "surrounding_text": surrounding_text[:500] if surrounding_text else "",
                })
            except:
                continue

# ============================================================================
# 7. Обработчик DOCX
# ============================================================================
class DocxProcessor:
    def __init__(self, path, source_model=""):
        self.path = Path(path)
        self.filename = self.path.name
        self.source_model = normalize_model_key(source_model)
        self.text_blocks = []
        self.figure_raw = []
        self.ocr_image_count = 0
        self.processing_time = 0

    def process(self):
        start_time = time.time()
        try:
            doc = DocxDocument(str(self.path))
        except Exception as e:
            print(f"  ❌ Error opening {self.filename}: {e}")
            return ([], [])

        print(f"  📝 {self.filename}")

        if self.source_model == "unknown":
            sample_text = "\n".join([p.text for p in doc.paragraphs[:20]])
            detected_model, conf = detect_model_in_text(sample_text)
            if conf > 0.1:
                self.source_model = detected_model
                print(f"    🏷️ Detected model: {MODEL_DISPLAY_NAMES.get(self.source_model, self.source_model)} (conf={conf})")

        self._process_paragraphs(doc)
        self._process_tables(doc)
        self._process_images(doc)

        self.processing_time = time.time() - start_time
        print(f"    ✅ Extracted {len(self.text_blocks)} blocks, {len(self.figure_raw)} figures, OCR images: {self.ocr_image_count} ({self.processing_time:.1f}s)")
        return (self.text_blocks, self.figure_raw)

    def _process_paragraphs(self, doc):
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else ""
            is_heading = False
            heading_level = 0
            if style_name and ("Heading" in style_name or "heading" in style_name or "Заголовок" in style_name):
                is_heading = True
                m = re.search(r'(\d+)', style_name)
                heading_level = int(m.group(1)) if m else 1
            block_type = "heading" if is_heading else "text"
            self.text_blocks.append({
                "text": text,
                "block_type": block_type,
                "heading_level": heading_level if is_heading else 0,
                "page": 0,
                "model": self.source_model,
                "source": self.filename,
            })

    def _process_tables(self, doc):
        for table_idx, table in enumerate(doc.tables):
            rows_data = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                rows_data.append(row_text)
            if not rows_data:
                continue
            table_text = "[ТАБЛИЦА]\n" + "\n".join(" | ".join(row) for row in rows_data) + "\n[/ТАБЛИЦА]"
            caption = find_figure_caption(table_text)
            self.text_blocks.append({
                "text": clean_text(table_text),
                "block_type": "table",
                "page": 0,
                "model": self.source_model,
                "source": self.filename,
                "caption": caption,
                "table_index": table_idx,
            })

    def _process_images(self, doc):
        image_parts = []
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_data = rel.target_part.blob
                    image_parts.append((rel.target_part.content_type, image_data))
                except:
                    continue
        if not image_parts:
            return
        for img_idx, (content_type, img_bytes) in enumerate(image_parts):
            try:
                ext_map = {"image/png": "png", "image/jpeg": "jpg", "image/jpg": "jpg",
                           "image/bmp": "bmp", "image/tiff": "tiff", "image/gif": "gif"}
                img_ext = ext_map.get(content_type, "png")
                if is_page_screenshot(img_bytes, min_size=(800, 1000)):
                    img = Image.open(io.BytesIO(img_bytes))
                    ocr_text = ocr_image(img)
                    if ocr_text:
                        self.text_blocks.append({
                            "text": ocr_text,
                            "block_type": "ocr_text",
                            "page": 0,
                            "model": self.source_model,
                            "source": self.filename,
                        })
                        self.ocr_image_count += 1
                    safe_name = re.sub(r'[^\w\-.]', '_', self.filename)
                    ocr_path = OCR_RESULTS_DIR / f"{safe_name}_docx_img{img_idx:03d}.txt"
                    ocr_path.write_text(ocr_text or "", encoding='utf-8')
                else:
                    img = Image.open(io.BytesIO(img_bytes))
                    w, h = img.size
                    if w < MIN_IMAGE_WIDTH or h < MIN_IMAGE_HEIGHT:
                        continue
                    surrounding = ""
                    for block in self.text_blocks[-5:]:
                        if block["source"] == self.filename:
                            surrounding += block["text"] + "\n"
                    caption = find_figure_caption(surrounding)
                    fig_id = f"{self.filename}_img{img_idx}"
                    self.figure_raw.append({
                        "fig_id": fig_id,
                        "source": self.filename,
                        "page": 0,
                        "model": self.source_model,
                        "width": w,
                        "height": h,
                        "ext": img_ext,
                        "caption": caption,
                        "img_bytes": img_bytes,
                        "surrounding_text": surrounding[:500],
                    })
            except:
                continue

# ============================================================================
# 8. Генератор source_models.json
# ============================================================================
def generate_source_models(docs_dir, output_dir):
    docs_dir = Path(docs_dir)
    source_models = {}
    pdf_files = sorted(docs_dir.rglob("*.pdf"))
    docx_files = sorted(docs_dir.rglob("*.docx"))
    all_files = [(f, "pdf") for f in pdf_files] + [(f, "docx") for f in docx_files]
    print(f"🔍 Analyzing {len(all_files)} files...")
    for filepath, ftype in all_files:
        filename = filepath.name
        model = "unknown"
        confidence = 0.0
        method = ""
        fn_model = detect_model_from_filename(filename)
        if fn_model != "unknown":
            model = fn_model
            confidence = 0.8
            method = "filename"
        if confidence < 0.7:
            try:
                if ftype == "pdf":
                    doc = fitz.open(str(filepath))
                    sample_text = ""
                    for pn in range(min(5, len(doc))):
                        sample_text += doc[pn].get_text("text") + "\n"
                    doc.close()
                else:
                    doc = DocxDocument(str(filepath))
                    sample_text = "\n".join([p.text for p in doc.paragraphs[:20]])
                content_model, content_conf = detect_model_in_text(sample_text)
                if content_conf > confidence:
                    model = content_model
                    confidence = content_conf
                    method = "content"
            except:
                pass
        source_models[filename] = model
        status = "✅" if model != "unknown" else "❓"
        print(f"  {status} {filename[:50]:50s} → {model:12s} (via={method})")
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    sm_path = output_dir / "source_models.json"
    with open(sm_path, 'w', encoding='utf-8') as f:
        json.dump(source_models, f, ensure_ascii=False, indent=2)
    print(f"\n✅ source_models.json saved to {sm_path}")
    return source_models

# ============================================================================
# 9. Главный пайплайн (только предобработка)
# ============================================================================
class KoibPreprocessingPipeline:
    def __init__(self, docs_dir, output_dir, source_models=None):
        self.docs_dir = Path(docs_dir)
        self.output_dir = Path(output_dir)
        self.source_models = source_models or {}
        self.text_blocks = []
        self.figures_index = []
        self.processing_log = []
        self.classification_report = {}
        self.total_ocr_pages = 0
        self.total_files = 0

    def _get_source_model(self, filename):
        if filename in self.source_models:
            return self.source_models[filename]
        return detect_model_from_filename(filename)

    def _save_figure(self, img_bytes, img_ext, fig_id, model, caption, page, source, surrounding=""):
        model_dir = FIGURES_DIR / model
        model_dir.mkdir(parents=True, exist_ok=True)
        safe_fig_id = re.sub(r'[^\w\-.]', '_', fig_id)
        fig_path = model_dir / f"{safe_fig_id}.{img_ext}"
        try:
            fig_path.write_bytes(img_bytes)
        except:
            return
        self.figures_index.append({
            "fig_id": fig_id,
            "fig_path": str(fig_path),
            "model": model,
            "source": source,
            "page": page,
            "caption": caption,
            "surrounding_text": surrounding,
        })

    def _process_pdf_figures(self, candidates):
        for fig in candidates:
            if fig.get("caption") or len(fig.get("surrounding_text", "")) > 50:
                self._save_figure(
                    img_bytes=fig["img_bytes"],
                    img_ext=fig["ext"],
                    fig_id=fig["fig_id"],
                    model=fig.get("model", "unknown"),
                    caption=fig.get("caption", ""),
                    page=fig.get("page", 0),
                    source=fig["source"],
                    surrounding=fig.get("surrounding_text", ""),
                )

    def _process_docx_figures(self, figure_raw):
        for fig in figure_raw:
            if fig.get("caption") or len(fig.get("surrounding_text", "")) > 50:
                self._save_figure(
                    img_bytes=fig["img_bytes"],
                    img_ext=fig["ext"],
                    fig_id=fig["fig_id"],
                    model=fig.get("model", "unknown"),
                    caption=fig.get("caption", ""),
                    page=fig.get("page", 0),
                    source=fig["source"],
                    surrounding=fig.get("surrounding_text", ""),
                )

    def process_all(self):
        print("="*70)
        print("🚀 KOIB Preprocessing Pipeline — extracting text and figures")
        print("="*70)
        pdf_files = sorted(self.docs_dir.rglob("*.pdf"))
        docx_files = sorted(self.docs_dir.rglob("*.docx"))
        self.total_files = len(pdf_files) + len(docx_files)
        if self.total_files == 0:
            print("❌ No PDF/DOCX files found!")
            return []
        print(f"📁 Found {len(pdf_files)} PDF, {len(docx_files)} DOCX")

        # PDFs
        print("\n--- Processing PDFs ---")
        for pdf_path in pdf_files:
            model = self._get_source_model(pdf_path.name)
            print(f"\n  Processing: {pdf_path.name} [model={model}]")
            processor = PdfProcessor(pdf_path, source_model=model)
            blocks, figures = processor.process()
            self.text_blocks.extend(blocks)
            self.total_ocr_pages += processor.ocr_page_count
            self._process_pdf_figures(figures)
            self.processing_log.append({
                "file": pdf_path.name,
                "type": "pdf",
                "model": processor.source_model,
                "text_blocks": len(blocks),
                "figures": len(figures),
                "ocr_pages": processor.ocr_page_count,
                "time_sec": round(processor.processing_time, 2),
            })

        # DOCXs
        print("\n--- Processing DOCXs ---")
        for docx_path in docx_files:
            model = self._get_source_model(docx_path.name)
            print(f"\n  Processing: {docx_path.name} [model={model}]")
            processor = DocxProcessor(docx_path, source_model=model)
            blocks, figure_raw = processor.process()
            self.text_blocks.extend(blocks)
            self.total_ocr_pages += processor.ocr_image_count
            self._process_docx_figures(figure_raw)
            self.processing_log.append({
                "file": docx_path.name,
                "type": "docx",
                "model": processor.source_model,
                "text_blocks": len(blocks),
                "figures": len(figure_raw),
                "ocr_pages": processor.ocr_image_count,
                "time_sec": round(processor.processing_time, 2),
            })

        self._build_classification_report()
        self.save_artifacts()
        return self.text_blocks

    def _build_classification_report(self):
        report = {
            "total_files": self.total_files,
            "by_type": defaultdict(int),
            "by_model": defaultdict(int),
            "by_model_type": defaultdict(int),
            "total_text_blocks": len(self.text_blocks),
            "total_figures": len(self.figures_index),
            "total_ocr_pages": self.total_ocr_pages,
            "files": [],
        }
        for entry in self.processing_log:
            report["by_type"][entry["type"]] += 1
            report["by_model"][entry["model"]] += 1
            report["by_model_type"][f"{entry['model']}_{entry['type']}"] += 1
            report["files"].append(entry)
        self.classification_report = dict(report)

    def save_artifacts(self):
        print("\n--- Saving artifacts ---")
        # Сохраняем text_blocks.json
        blocks_path = METADATA_DIR / "text_blocks.json"
        with open(blocks_path, 'w', encoding='utf-8') as f:
            json.dump(self.text_blocks, f, ensure_ascii=False, indent=2)
        print(f"  ✅ Text blocks saved: {blocks_path} ({len(self.text_blocks)} blocks)")

        # Сохраняем figures_index.json
        figures_path = METADATA_DIR / "figures_index.json"
        with open(figures_path, 'w', encoding='utf-8') as f:
            json.dump(self.figures_index, f, ensure_ascii=False, indent=2)
        print(f"  ✅ Figures index saved: {figures_path} ({len(self.figures_index)} figures)")

        # Сохраняем processing_log.json
        log_path = LOGS_DIR / "processing_log.json"
        with open(log_path, 'w', encoding='utf-8') as f:
            json.dump(self.processing_log, f, ensure_ascii=False, indent=2)
        print(f"  ✅ Processing log saved: {log_path}")

        # Сохраняем classification_report.csv
        csv_path = CLASSIFIED_DIR / "classification_report.csv"
        csv_path.parent.mkdir(parents=True, exist_ok=True)
        with open(csv_path, 'w', encoding='utf-8-sig', newline='') as f:
            writer = csv.writer(f)
            writer.writerow(["File", "Type", "Model", "Text Blocks", "Figures", "OCR Pages", "Time (s)"])
            for entry in self.processing_log:
                writer.writerow([entry["file"], entry["type"], entry["model"],
                                 entry["text_blocks"], entry["figures"],
                                 entry["ocr_pages"], entry["time_sec"]])
        print(f"  ✅ Classification report saved: {csv_path}")

    def print_summary(self):
        cr = self.classification_report
        print("\n" + "="*70)
        print("📊 PREPROCESSING SUMMARY")
        print("="*70)
        print(f"  Total files:          {cr.get('total_files',0)}")
        print(f"  By type: PDF={cr['by_type'].get('pdf',0)} DOCX={cr['by_type'].get('docx',0)}")
        print(f"  By model: {dict(cr.get('by_model',{}))}")
        print(f"  Text blocks:          {cr.get('total_text_blocks',0)}")
        print(f"  Figures extracted:    {cr.get('total_figures',0)}")
        print(f"  OCR pages/images:     {cr.get('total_ocr_pages',0)}")
        print("="*70)

# ============================================================================
# 10. MAIN – запуск предобработки
# ============================================================================
def main():
    print("╔" + "═"*68 + "╗")
    print("║" + "  KOIB RAG v4 – PART 1: PREPROCESSING".center(68) + "║")
    print("╚" + "═"*68 + "╝")
    start_time = time.time()

    # Генерация source_models.json (если ещё нет)
    sm_path = METADATA_DIR / "source_models.json"
    if sm_path.exists():
        print(f"\n📂 Loading existing source_models.json from {sm_path}")
        with open(sm_path, 'r', encoding='utf-8') as f:
            source_models = json.load(f)
    else:
        print("\n📂 Generating source_models.json...")
        source_models = generate_source_models(DOCS_DIR, METADATA_DIR)

    # Запуск пайплайна предобработки
    pipeline = KoibPreprocessingPipeline(DOCS_DIR, OUTPUT_DIR, source_models)
    text_blocks = pipeline.process_all()
    pipeline.save_artifacts()
    pipeline.print_summary()

    elapsed = time.time() - start_time
    print(f"\n⏱️ Total preprocessing time: {elapsed:.1f}s ({elapsed/60:.1f} min)")
    print("✅ PART 1 FINISHED. Now you can run PART 2 (index building).")

if __name__ == "__main__":
    main()